import os
from typing import Tuple, Dict, Any, List
from pypdf import PdfReader
import docx

class DocumentParserService:
    def parse_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extracts text from a PDF file.
        Returns a tuple of (extracted_text, metadata).
        """
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
                
        extracted_text = "\n".join(text_parts)
        
        metadata = {
            "pages_count": len(reader.pages),
            "author": reader.metadata.author if reader.metadata else None,
            "creator": reader.metadata.creator if reader.metadata else None,
            "producer": reader.metadata.producer if reader.metadata else None,
            "subject": reader.metadata.subject if reader.metadata else None,
            "title": reader.metadata.title if reader.metadata else None,
        }
        # Remove keys with None values
        metadata = {k: v for k, v in metadata.items() if v is not None}
        
        return extracted_text, metadata

    def parse_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extracts text and tables from a Word (.docx) file.
        Returns a tuple of (extracted_text, metadata).
        """
        doc = docx.Document(file_path)
        
        # Extract paragraphs text
        paragraphs_text = [p.text for p in doc.paragraphs if p.text]
        extracted_text = "\n".join(paragraphs_text)
        
        # Extract tables
        tables_data = []
        for t_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_cells)
            tables_data.append({
                "table_index": t_idx,
                "rows": table_rows
            })
            
        metadata = {
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
        }
        if tables_data:
            metadata["tables"] = tables_data
            
        return extracted_text, metadata

    def extract_text_and_metadata(self, file_path: str, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Identifies file extension and routes it to the appropriate parser.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = os.path.splitext(filename.lower())[1]
        
        if ext == ".pdf":
            return self.parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self.parse_docx(file_path)
        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return text, {"file_type": "text"}
        else:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                return text, {"file_type": "unknown_text_fallback"}
            except Exception:
                return "", {"file_type": "unsupported"}

document_parser = DocumentParserService()
