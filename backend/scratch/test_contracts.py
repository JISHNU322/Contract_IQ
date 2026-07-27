import os
import time
import uuid
import requests
import docx

BASE_URL = "http://127.0.0.1:8000"

def create_sample_docx(filename: str):
    doc = docx.Document()
    doc.add_heading("Contract Agreement", level=0)
    doc.add_paragraph("This agreement is made between Party A and Party B.")
    doc.add_paragraph("Term: 12 months from the effective date.")
    
    # Add a table
    table = doc.add_table(rows=3, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Parameter"
    hdr_cells[1].text = "Value"
    
    row_cells1 = table.rows[1].cells
    row_cells1[0].text = "Effective Date"
    row_cells1[1].text = "2026-08-01"
    
    row_cells2 = table.rows[2].cells
    row_cells2[0].text = "Value"
    row_cells2[1].text = "$50,000"
    
    doc.save(filename)
    print(f"Created sample DOCX: {filename}")

def run_tests():
    # 1. Create a unique user to avoid conflict
    email = f"test_{uuid.uuid4().hex[:6]}@example.com"
    password = "testpassword123"
    
    print(f"Registering user: {email}...")
    register_response = requests.post(
        f"{BASE_URL}/api/auth/register",
        json={"email": email, "password": password, "full_name": "Test User"}
    )
    assert register_response.status_code == 200, f"Register failed: {register_response.text}"
    user_data = register_response.json()
    print(f"Registered user ID: {user_data['id']}")

    # 2. Login to get token
    print("Logging in...")
    login_response = requests.post(
        f"{BASE_URL}/api/auth/login",
        data={"username": email, "password": password}
    )
    assert login_response.status_code == 200, f"Login failed: {login_response.text}"
    token = login_response.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    print("Login successful.")

    # 3. Create a sample contract docx
    docx_filename = f"sample_contract_{uuid.uuid4().hex[:6]}.docx"
    create_sample_docx(docx_filename)

    try:
        # 4. Upload contract
        print("Uploading contract...")
        with open(docx_filename, "rb") as f:
            upload_response = requests.post(
                f"{BASE_URL}/api/contracts/upload",
                headers=headers,
                files={"file": (docx_filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        assert upload_response.status_code == 201, f"Upload failed: {upload_response.text}"
        contract_data = upload_response.json()
        contract_id = contract_data["id"]
        print(f"Uploaded successfully. Contract ID: {contract_id}, Status: {contract_data['status']}")

        # 5. Poll status until 'parsed' or 'failed'
        max_attempts = 10
        status = contract_data["status"]
        for attempt in range(max_attempts):
            print(f"Checking status (Attempt {attempt+1}/{max_attempts})...")
            get_response = requests.get(
                f"{BASE_URL}/api/contracts/{contract_id}",
                headers=headers
            )
            assert get_response.status_code == 200, f"Get contract failed: {get_response.text}"
            contract_data = get_response.json()
            status = contract_data["status"]
            print(f"Current Status: {status}")
            if status in ["parsed", "failed"]:
                break
            time.sleep(1)

        assert status == "parsed", f"Contract parsing failed, final state: {status}. Metadata: {contract_data.get('extracted_metadata')}"
        print("Parsing verification passed!")
        print(f"Extracted plain text snippet: {repr(contract_data['parsed_text'][:100])}...")
        print(f"Extracted metadata: {contract_data['extracted_metadata']}")

        # 6. Verify table parsing in metadata
        metadata = contract_data["extracted_metadata"]
        assert "tables" in metadata, "No tables found in metadata!"
        tables = metadata["tables"]
        assert len(tables) > 0, "Tables list is empty"
        rows = tables[0]["rows"]
        assert rows[1][0] == "Effective Date", f"Expected 'Effective Date', got: {rows[1][0]}"
        print("Table parsing verification passed!")

        # 7. Download original file
        print("Downloading file...")
        download_response = requests.get(
            f"{BASE_URL}/api/contracts/{contract_id}/download",
            headers=headers
        )
        assert download_response.status_code == 200, f"Download failed: {download_response.text}"
        assert len(download_response.content) == os.path.getsize(docx_filename), "Downloaded file size mismatch"
        print("Download verification passed!")

        # 8. Delete contract
        print("Deleting contract...")
        delete_response = requests.delete(
            f"{BASE_URL}/api/contracts/{contract_id}",
            headers=headers
        )
        assert delete_response.status_code == 204, f"Delete failed: {delete_response.text}"
        print("Delete contract API verification passed!")

        # 9. Verify deletion from list
        list_response = requests.get(
            f"{BASE_URL}/api/contracts/",
            headers=headers
        )
        assert list_response.status_code == 200, f"List failed: {list_response.text}"
        contracts_list = list_response.json()
        assert all(c["id"] != contract_id for c in contracts_list), "Contract was not removed from list"
        print("List verification after deletion passed!")

        print("ALL TESTS PASSED SUCCESSFULLY!")

    finally:
        # Cleanup local file
        if os.path.exists(docx_filename):
            os.remove(docx_filename)

if __name__ == "__main__":
    run_tests()
